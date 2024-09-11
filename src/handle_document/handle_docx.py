from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import zipfile
import re

import win32com.client

class WordHandler():
    def __init__(self):
        pass

    @staticmethod
    def extract_page_num(file_name):
        escaped_path = file_name.replace("/", "\\")
        page_num = 0
        try:
            if escaped_path.endswith('.docx'):
                with zipfile.ZipFile(escaped_path) as docx_object:
                    docx_property_file_data = docx_object.read('docProps/app.xml').decode()
                    page_num = re.search(r"<Pages>(\d+)</Pages>", docx_property_file_data).group(1)
            elif escaped_path.endswith('.doc'):
                app = win32com.client.Dispatch('Word.Application')
                doc = app.Documents.Open(escaped_path)
                page_num = len(doc.ActiveWindow.ActivePane.Pages)

                doc.Close(False)
                app.Quit()
        except Exception as e:
            print(f"Error extracting page number from {file_name}: {e}")

        print(f"{file_name}'s Page Number: {page_num}")
        return int(page_num)

    @staticmethod
    def create_cover_page(doc_name, pjt_no, dept_name, person_name, output_file):
        document = Document()

        # 문서 제목 추가
        document.add_heading(doc_name, 0)

        # 빈 줄 추가
        document.add_paragraph('')

        # 프로젝트 넘버 추가
        department = document.add_paragraph()
        department_run = department.add_run('Project No: ')
        department_run.bold = True
        department.add_run(pjt_no)

        # 담당 부서 추가
        department = document.add_paragraph()
        department_run = department.add_run('담당 부서: ')
        department_run.bold = True
        department.add_run(dept_name)

        # 담당자 추가
        person = document.add_paragraph()
        person_run = person.add_run('담당자: ')
        person_run.bold = True
        person.add_run(person_name)

        # 페이지 여백 설정
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        document.save(output_file)

    @staticmethod
    def create_index_page(files_description, files_pages, output_file):
        'example'
        # files_pages = (2, 3, 1)
        # files_description = ('first file', 'second file', 'third file')

        num_files = len(files_description)
        document = Document()

        document.add_heading('INDEX', 0)

        # 페이지 여백 설정
        document.sections[0].left_margin = Cm(2.6)
        document.sections[0].right_margin = Cm(2)
        document.sections[0].top_margin = Cm(2)
        document.sections[0].bottom_margin = Cm(2)

        table = document.add_table(rows=num_files + 1, cols=4)
        table.style = document.styles['Table Grid']
        heading_cells = table.rows[0].cells
        heading_text = ('No', 'DESCRIPTION', 'PAGE', 'REMARKS')
        heading_width = (1.19, 15.5, 2.25, 3.45)

        # heading cell
        for i, (text, width) in enumerate(zip(heading_text, heading_width)):
            heading_cells[i].text = text
            heading_cells[i].width = Cm(width)
            heading_cells[i].paragraphs[0].style.font.name = '맑은 고딕'
            if i > 0:
                heading_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # content cell
        row = table.rows[1]
        total_page_num = 3 # title page + index page + first page + ..
        for i in range(1, num_files + 1):
            row = table.rows[i]
            row.cells[0].text = f'{i}'
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[1].text = files_description[i - 1]
            if total_page_num < 10 and total_page_num + int(files_pages[i - 1]) < 10:
                row.cells[2].text = f'0{total_page_num}~0{total_page_num + int(files_pages[i - 1]-1)}'
            else:
                row.cells[2].text = f'{total_page_num:02}~{total_page_num + int(files_pages[i - 1]-1):02}'

            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            total_page_num = total_page_num + int(files_pages[i - 1])
        document.save(output_file)

if __name__ == '__main__':
    page_num = WordHandler.get_page_num("소프트웨어 설계서.docx")
    print(page_num)